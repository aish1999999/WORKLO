#!/usr/bin/env python3
"""
Resume TXT Converter - Enhanced with Strong Bullet Point Generation

Includes professional bullet point generation prompt for achievement-focused,
quantifiable, and ATS-optimized resume content.
"""

import re
import logging
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

logger = logging.getLogger(__name__)

class ResumeTextConverter:
    """Enhanced resume converter with strong bullet point generation."""

    def __init__(self):
        self.section_markers = {
            'HEADER': '===HEADER===',
            'EDUCATION': '===EDUCATION===',
             'PROJECTS': '===PROJECTS===',
            'EXPERIENCE': '===PROFESSIONAL EXPERIENCE===',
            'SKILLS': '===TECHNICAL SKILLS===',
        }

        self.format_specs = {
            'font_name': 'Times New Roman',
            'name_size': 16,                  # Name: 16pt, BOLD
            'contact_size': 11,               # Contact: 11pt, regular
            'section_header_size': 13,        # Section headers: 13pt, BOLD + UNDERLINE
            'job_title_size': 12,             # Job titles: 12pt, BOLD
            'body_size': 11,                  # Bullet points & body: 11pt, regular
            'left_margin': 0.5,               # 0.5" left margin
            'right_margin': 0.5,              # 0.5" right margin
            'top_margin': 0.3,                # 0.3" top margin (ultra-narrow for 1-page fit)
            'bottom_margin': 0.3,             # 0.3" bottom margin (ultra-narrow for 1-page fit)
            'no_spacing': 0,                  # 0pt between bullets
        }

    def docx_to_structured_txt(self, docx_path: str) -> str:
        """Convert DOCX resume to structured TXT format with section markers."""
        try:
            doc = Document(docx_path)
            structured_lines = []
            current_section = None

            # Process paragraphs in the exact order they appear
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect section headers (ALL CAPS or specific keywords)
                if self._is_section_header(text):
                    section_key = self._map_to_section_key(text)
                    if section_key:
                        structured_lines.append(f"\n{self.section_markers[section_key]}")
                        current_section = section_key
                        structured_lines.append(text)
                    else:
                        structured_lines.append(text)
                else:
                    # Regular content
                    if not current_section:
                        # Before first section - treat as header
                        if not any(marker in '\n'.join(structured_lines) for marker in self.section_markers.values()):
                            structured_lines.insert(0, self.section_markers['HEADER'])
                            current_section = 'HEADER'

                    structured_lines.append(text)

            result = '\n'.join(structured_lines)
            logger.info(f"Converted DOCX to structured TXT ({len(structured_lines)} lines)")
            return result

        except Exception as e:
            logger.error(f"Error converting DOCX to TXT: {e}")
            return ""

    def _is_section_header(self, text: str) -> bool:
        """Check if text is a section header."""
        text_upper = text.upper().strip()

        # Check for specific section keywords
        section_keywords = [
            'EDUCATION', 'PROFESSIONAL EXPERIENCE', 'EXPERIENCE', 'TECHNICAL SKILLS',
            'SKILLS', 'CORE SKILLS', 'SUMMARY', 'PROJECTS'
        ]

        return any(keyword in text_upper for keyword in section_keywords)

    def _map_to_section_key(self, header_text: str) -> str:
        """Map header text to section key."""
        header_upper = header_text.upper()

        if 'EDUCATION' in header_upper:
            return 'EDUCATION'
        elif 'PROJECTS' in header_upper:
            return 'PROJECTS'
        elif 'EXPERIENCE' in header_upper:
            return 'EXPERIENCE'
        elif 'SKILL' in header_upper or 'TECHNICAL' in header_upper:
            return 'SKILLS'
      

        return None

    def structured_txt_to_docx(self, structured_txt: str, master_docx_path: str, output_path: str) -> bool:
        """Convert structured TXT back to DOCX with exact formatting."""
        try:
            # Clean the structured text first
            cleaned_txt = self._clean_structured_txt(structured_txt)

            # Parse the cleaned structured text
            sections = self._parse_ai_structured_txt(cleaned_txt)

            # Create new document with exact formatting
            new_doc = Document()
            self._set_document_margins(new_doc)

            # Add each section with proper formatting
            self._add_header_section(new_doc, sections.get('HEADER', []))
            self._add_education_section(new_doc, sections.get('EDUCATION', []))
            self._add_projects_section(new_doc, sections.get('PROJECTS', []))
            self._add_experience_section(new_doc, sections.get('EXPERIENCE', []))
            self._add_skills_section(new_doc, sections.get('SKILLS', []))

            # Save the formatted document
            new_doc.save(output_path)
            logger.info(f"Created properly formatted resume: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Error creating formatted DOCX: {e}")
            return False

    def _clean_structured_txt(self, structured_txt: str) -> str:
        """Clean the structured text to remove formatting instructions."""
        lines = structured_txt.split('\n')
        cleaned_lines = []
        skip_formatting_section = False

        for line in lines:
            line = line.strip()

            # Skip formatting instruction sections
            if any(keyword in line for keyword in ['**FORMATTING SPECIFICATIONS:**', '**CONTENT REQUIREMENTS:**', '**WHAT NEVER CHANGES:**', '**Instructions:**', '**Output Format:**']):
                skip_formatting_section = True
                continue

            # Check if we're still in a formatting section
            if skip_formatting_section:
                if line.startswith('===') or (line and not line.startswith('-') and not line.startswith('**') and not line.startswith('*') and not line.startswith('1.') and not line.startswith('2.') and not line.startswith('3.') and not line.startswith('4.') and not line.startswith('5.')):
                    skip_formatting_section = False
                else:
                    continue

            # Keep the line if we're not skipping
            if not skip_formatting_section:
                cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    def _parse_ai_structured_txt(self, structured_txt: str) -> Dict[str, List[str]]:
        """Parse AI-generated structured text into sections."""
        sections = {}
        current_section = None
        current_lines = []

        for line in structured_txt.split('\n'):
            line = line.strip()
            if not line:
                continue

            # Check for section markers
            if '===HEADER===' in line:
                current_section = self._save_and_start_section(sections, current_section, current_lines, 'HEADER')
                current_lines = []
            elif '===EDUCATION===' in line:
                current_section = self._save_and_start_section(sections, current_section, current_lines, 'EDUCATION')
                current_lines = []
            elif '===PROJECTS===' in line:
                current_section = self._save_and_start_section(sections, current_section, current_lines, 'PROJECTS')
                current_lines = []
            elif '===PROFESSIONAL EXPERIENCE===' in line or '===EXPERIENCE===' in line:
                current_section = self._save_and_start_section(sections, current_section, current_lines, 'EXPERIENCE')
                current_lines = []
            elif '===TECHNICAL SKILLS===' in line or '===SKILLS===' in line:
                current_section = self._save_and_start_section(sections, current_section, current_lines, 'SKILLS')
                current_lines = []
            else:
                # Add content to current section
                if current_section:
                    current_lines.append(line)
                else:
                    # No section found yet, treat as header
                    if 'HEADER' not in sections:
                        sections['HEADER'] = []
                    sections['HEADER'].append(line)

        # Save final section
        if current_section and current_lines:
            sections[current_section] = current_lines

        logger.info(f"Parsed AI output into sections: {list(sections.keys())}")
        return sections

    def _save_and_start_section(self, sections: Dict, current_section: str, current_lines: List[str], new_section: str) -> str:
        """Helper to save current section and start new one."""
        if current_section and current_lines:
            sections[current_section] = current_lines
        return new_section

    def _set_document_margins(self, doc: Document) -> None:
        """Set margins: 0.5" left/right, 1.0" top/bottom."""
        section = doc.sections[0]
        section.left_margin = Inches(self.format_specs['left_margin'])
        section.right_margin = Inches(self.format_specs['right_margin'])
        section.top_margin = Inches(self.format_specs['top_margin'])
        section.bottom_margin = Inches(self.format_specs['bottom_margin'])

    def _add_header_section(self, doc: Document, header_lines: List[str]) -> None:
        """Add header: name and contact info (both centered)."""
        for i, line in enumerate(header_lines):
            if not line.strip():
                continue

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # CENTER ALL HEADER CONTENT

            if i == 0:  # Name - bold
                formatted_name = self._format_name_correctly(line)
                run = para.add_run(formatted_name)
                run.font.name = self.format_specs['font_name']
                run.font.size = Pt(self.format_specs['name_size'])
                run.bold = True
            else:  # Contact info - regular weight, still centered
                if self._is_contact_info(line):
                    run = para.add_run(line.strip())
                    run.font.name = self.format_specs['font_name']
                    run.font.size = Pt(self.format_specs['contact_size'])
                    run.bold = False
                else:
                    # Skip non-contact lines (like job titles)
                    continue

            para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
            para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

    def _add_education_section(self, doc: Document, education_lines: List[str]) -> None:
        """Add EDUCATION section with underlined header."""
        if not education_lines:
            return

        # Add section header: BOLD + UNDERLINED
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run("EDUCATION")
        run.font.name = self.format_specs['font_name']
        run.font.size = Pt(self.format_specs['section_header_size'])
        run.bold = True
        run.underline = True
        para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
        para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

        # Add education entries
        for line in education_lines:
            if line.strip():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(line.strip())
                run.font.name = self.format_specs['font_name']
                run.font.size = Pt(self.format_specs['job_title_size'])
                run.bold = False
                para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
                para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

    def _add_projects_section(self, doc: Document, projects_lines: List[str]) -> None:
        """Add PROJECTS with underlined header and bold project entries."""
        if not projects_lines:
            return

        # Add blank line before section header
        blank_para = doc.add_paragraph()
        blank_para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
        blank_para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

        # Add section header: BOLD + UNDERLINED
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run("PROJECTS")
        run.font.name = self.format_specs['font_name']
        run.font.size = Pt(self.format_specs['section_header_size'])
        run.bold = True
        run.underline = True
        para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
        para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

        # Add projects content
        for line in projects_lines:
            if not line.strip():
                continue

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if '|' in line:
                # Project entry line: Project Name | Organization | Dates - ALL BOLD
                run = para.add_run(line.strip())
                run.font.name = self.format_specs['font_name']
                run.font.size = Pt(self.format_specs['job_title_size'])
                run.bold = True
            else:
                # Bullet point line - add period at end
                content = line.strip()
                if not content.startswith('•'):
                    content = f"• {content}"

                # Add period if not already present
                if not content.endswith('.') and not content.endswith('?') and not content.endswith('!'):
                    content += '.'

                run = para.add_run(content)
                run.font.name = self.format_specs['font_name']
                run.font.size = Pt(self.format_specs['body_size'])
                run.bold = False

            para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
            para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])


    def _add_experience_section(self, doc: Document, experience_lines: List[str]) -> None:
        """Add PROFESSIONAL EXPERIENCE with underlined header and bold job entries."""
        if not experience_lines:
            return

        # Add blank line before section header
        blank_para = doc.add_paragraph()
        blank_para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
        blank_para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

        # Add section header: BOLD + UNDERLINED
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run("PROFESSIONAL EXPERIENCE")
        run.font.name = self.format_specs['font_name']
        run.font.size = Pt(self.format_specs['section_header_size'])
        run.bold = True
        run.underline = True
        para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
        para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

        # Add experience content
        for line in experience_lines:
            if not line.strip():
                continue

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if '|' in line:
                # Job entry line: Company | Role | Location | Duration - ALL BOLD
                run = para.add_run(line.strip())
                run.font.name = self.format_specs['font_name']
                run.font.size = Pt(self.format_specs['job_title_size'])
                run.bold = True
            else:
                # Bullet point line - add period at end
                content = line.strip()
                if not content.startswith('•'):
                    content = f"• {content}"

                # Add period if not already present
                if not content.endswith('.') and not content.endswith('?') and not content.endswith('!'):
                    content += '.'

                run = para.add_run(content)
                run.font.name = self.format_specs['font_name']
                run.font.size = Pt(self.format_specs['body_size'])
                run.bold = False

            para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
            para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

    def _add_skills_section(self, doc: Document, skills_lines: List[str]) -> None:
        """Add TECHNICAL SKILLS with underlined header and clean formatting."""
        if not skills_lines:
            return

        # Add blank line before section header
        blank_para = doc.add_paragraph()
        blank_para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
        blank_para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

        # Add section header: BOLD + UNDERLINED
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run("TECHNICAL SKILLS")
        run.font.name = self.format_specs['font_name']
        run.font.size = Pt(self.format_specs['section_header_size'])
        run.bold = True
        run.underline = True
        para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
        para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

        # Add skills content
        for line in skills_lines:
            if not line.strip():
                continue

            # Skip any remaining formatting instruction lines
            if any(keyword in line for keyword in ['FORMATTING', 'CONTENT REQUIREMENTS', 'WHAT NEVER CHANGES', 'Instructions:', 'Action-Oriented:', 'STAR Flow:', 'Quantifiable Results:', 'Concise & Specific:', 'ATS Alignment:', 'Output Format:']):
                continue

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Parse skills line - handle both ** and plain formatting
            if ':' in line:
                # Split on first colon
                parts = line.split(':', 1)
                category_part = parts[0].strip()
                skills_part = parts[1].strip() if len(parts) > 1 else ""

                # Clean category (remove ** if present)
                category = category_part.replace('**', '').replace('*', '').strip()

                # Clean skills (remove ** if present)
                skills = skills_part.replace('**', '').replace('*', '').strip()

                # Add bold category
                run1 = para.add_run(f"{category}:")
                run1.font.name = self.format_specs['font_name']
                run1.font.size = Pt(self.format_specs['body_size'])
                run1.bold = True

                # Add regular skills
                if skills:
                    run2 = para.add_run(f" {skills}")
                    run2.font.name = self.format_specs['font_name']
                    run2.font.size = Pt(self.format_specs['body_size'])
                    run2.bold = False
            else:
                # Regular skills line - clean any ** formatting
                clean_line = line.replace('**', '').replace('*', '').strip()
                run = para.add_run(clean_line)
                run.font.name = self.format_specs['font_name']
                run.font.size = Pt(self.format_specs['body_size'])
                run.bold = False

            para.paragraph_format.space_before = Pt(self.format_specs['no_spacing'])
            para.paragraph_format.space_after = Pt(self.format_specs['no_spacing'])

    def _format_name_correctly(self, name_line: str) -> str:
        """Format name: only first letters of first/last name capitalized."""
        words = name_line.strip().split()
        name_words = []

        for word in words:
            if not any(char in word for char in ['@', '|', 'http', '.com', '832', '713']):
                name_words.append(word)
            else:
                break

        if len(name_words) >= 2:
            first_name = name_words[0]
            last_name = name_words[-1]

            formatted_first = first_name[0].upper() + first_name[1:].lower() if first_name else ""
            formatted_last = last_name[0].upper() + last_name[1:].lower() if last_name else ""

            return f"{formatted_first} {formatted_last}"

        return name_line.strip()

    def _is_contact_info(self, line: str) -> bool:
        """Check if line contains contact information."""
        line_lower = line.lower()
        return (
            '@' in line or
            'linkedin' in line_lower or
            'http' in line_lower or
            re.search(r'\d{3}[-\s]?\d{3}[-\s]?\d{4}', line) or
            '|' in line
        )


# ENHANCED SYSTEM PROMPT with Professional Bullet Point Generation
ENHANCED_SYSTEM_PROMPT = """You are an expert resume writer and professional bullet-point generator. Your task is to transform job descriptions into achievement-focused, quantifiable, and ATS-optimized resume content.

**CRITICAL**: Return content in this EXACT structured format:

===HEADER===
Aishwaryeshwar Manickavel
ayeshwarm@gmail.com | 832-775-2886 | www.linkedin.com/in/aishwarnick

===EDUCATION===
University of Houston | Master of Science in Industrial Engineering | Houston, TX, USA |  December 2025
PSG College of Technology | Bachelor of Engineering in Metallurgical Engineering | TN, India | July 2023

===PROJECTS===
Quality 4.0: Digital Twin with LLM-Powered SPC Analytics| University of Houston| Feb 2025 - Present
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]

Time Series Predictive Analytics-Operational Performance |University of Houston| Jan 2024-May 2024
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]

===PROFESSIONAL EXPERIENCE===
University of Houston | Operations Lead | Houston, TX | USA | May 2024 – Feb 2025
• Optimized workforce scheduling algorithms using demand forecasting and resource allocation strategies, improving staff utilization by 20% while maintaining service quality standards.
• Reduced operational downtime by 15% through implementation of preventive maintenance protocols and real-time equipment tracking systems, saving $8,000+ in potential revenue loss.
• Managed end-to-end event logistics for 25+ campus events per semester, ensuring 100% safety compliance and coordinating cross-functional teams of 15+ staff members.

Centre for Automotive Energy Materials at IIT Madras | [Tailored Job Role] | Chennai, TN | India | Jun 2022 -- December 2023
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]

A. V.C Pvt Ltd. | [Tailored Job Role] | TN, India | January 2021 -- May 2022
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]

===TECHNICAL SKILLS===
[JD Category 1]: [5-6 unique job-specific skills from JD, under 100 chars]
Data & Analytics: Power BI, Tableau, SQL, Excel, Python, Pandas, NumPy
[JD Category 2]: [5-6 unique job-specific skills from JD, under 100 chars, NO duplicates from Category 1]
Design & Simulation: CAD, SolidWorks, Creo, 3D Printing, Arena Simulation, Minitab
[JD Category 3]: [5-6 unique job-specific skills from JD, under 100 chars, NO duplicates from Categories 1 or 2]
Process & Quality: Lean Six Sigma, Root Cause Analysis, FMEA, SPC, JMP

**CRITICAL - ALTERNATING SKILLS STRUCTURE (EXACTLY 6 CATEGORIES):**
You MUST create EXACTLY 6 skill categories in a strict alternating pattern:
1. [JD Category 1] - From job description (5-6 UNIQUE skills, <100 chars total)
2. Data & Analytics (Default) - Power BI, Tableau, SQL, Excel, Python, Pandas, NumPy
3. [JD Category 2] - From job description (5-6 UNIQUE skills NOT in Category 1, <100 chars total)
4. Design & Simulation (Default) - CAD, SolidWorks, Creo, 3D Printing, Arena Simulation, Minitab
5. [JD Category 3] - From job description (5-6 UNIQUE skills NOT in Categories 1 or 3, <100 chars total)
6. Process & Quality (Default) - Lean Six Sigma, Root Cause Analysis, FMEA, SPC, JMP

**MANDATORY NO-DUPLICATION RULE:**
- Each JD-specific skill can appear in ONLY ONE category
- NO duplicate skills across all 6 categories

**BULLET POINT REQUIREMENTS:**

1. **Action-Oriented:** Start every bullet with a strong action verb (e.g., Negotiated, Implemented, Reduced, Streamlined, Optimized, Led, Developed, Managed).

2. **STAR Flow:** Each bullet must follow *Situation → Task → Action → Result*:
   * What you did
   * How you did it (method, tools, process)
   * Why it mattered (impact with measurable outcome)

3. **Quantifiable Results:** MUST include specific hard numbers appropriate for entry/mid-level roles:
   * Exact percentages: "18% cost reduction" not "significant savings"
   * Specific dollar amounts: "$85K savings" or "across $470K budget"
   * Precise quantities: "8+ contracts", "35+ vendors", "6 team members"
   * Before/after comparisons: "from 16 to 12.5 weeks", "4.2x to 5.8x turnover"
   * Entry-level metrics: budgets ($50K-$500K), teams (3-8 people), improvements (10-25%)
   * Mid-level metrics: budgets ($500K-$2M), teams (8-15 people), improvements (20-35%)

4. **Concise & Specific:** Keep bullets 1–2 lines max. Avoid vague or generic phrasing.

5. **ATS Alignment:** Use keywords and industry terms from the job description itself.

6. **Format:** **[Action] + [Method/Tool/Process] + [Quantified Result/Impact]**.

**CONTENT REQUIREMENTS:**
- Rewrite ALL content to match the job description and keywords
- Keep company names and university names UNCHANGED
- Job entry format: Company | Job Role | Location | Duration
- Use exactly 2 bullet points per project (Quality 4.0 and Time Series)
- For "Quality 4.0: Digital Twin" project - MUST mention Unity (simulation platform) and Python, focus on simulation modeling, real-time monitoring, and LLM integration
- For "Time Series Predictive Analytics" project - Focus on predictive analytics, machine learning, data visualization, and operational performance optimization
- Use exactly 3 bullet points for Centre for Automotive Energy Materials at IIT Madras
- Use exactly 3 bullet points for A.V.C Pvt Ltd
- Create EXACTLY 6 skill categories in alternating pattern (3 JD-specific + 3 default)
- Each skill category line must be under 100 characters
- Natural keyword integration throughout
- Name: "Aishwaryeshwar Manickavel" (first letters only capitalized)
- NO duplicate skills across all 6 categories

Return ONLY the structured text with ===SECTION=== markers above. Do not include any formatting specifications or instructions in your response."""


if __name__ == "__main__":
    print("Enhanced Resume Converter with Strong Bullet Points:")
    print("✓ Professional bullet-point generator integrated")
    print("✓ STAR method (Situation → Task → Action → Result)")
    print("✓ Action-oriented with strong verbs")
    print("✓ Quantifiable results and metrics")
    print("✓ ATS-optimized with job description keywords")
    print("✓ Concise and specific (1-2 lines max)")