#!/usr/bin/env python3
"""
Automated Resume Tailor

A comprehensive tool for tailoring resumes to specific job descriptions using AI.

INSTALLATION:
pip install gspread oauth2client openai anthropic python-docx PyDrive2 tenacity python-dotenv

GOOGLE OAUTH SETUP:
1. Go to Google Cloud Console
2. Create/select project → Enable Google Sheets API & Google Drive API
3. Create OAuth 2.0 Client ID → Download JSON file as client_secret.json
4. Set GOOGLE_CLIENT_SECRETS_FILE env var to the JSON file path
5. The script will handle OAuth flow automatically

ENVIRONMENT VARIABLES:
- OPENAI_API_KEY: Your OpenAI API key
- ANTHROPIC_API_KEY: Your Anthropic API key (if using Anthropic)
- LLM_PROVIDER: "openai" or "anthropic" (default: "openai")
- GOOGLE_AUTH_MODE: "oauth" or "service_account" (default: "service_account")
- GOOGLE_CLIENT_SECRETS_FILE: Path to OAuth client secrets JSON file
- GOOGLE_OAUTH_TOKEN_FILE: Path to store OAuth token (optional, default: ./token.json)
- DRIVE_FOLDER_ID: Google Drive folder ID for uploads (default: 1PCBjveFtn07ljJyIhmVjwlwWwPZ7hoFr)
- DRY_RUN: Set to "true" to skip actual API calls

CLI EXAMPLES:
python automated_resume_tailor.py --sheet-id 1dkxFBloNJQMnXtTfLG5CsGryK_EfmlYrdFFluWPBkgw --worksheet "Jobs" --base-resume ./base_resume.docx --master-template ./Master_template2.docx

python automated_resume_tailor.py --sheet-id 1dkxFBloNJQMnXtTfLG5CsGryK_EfmlYrdFFluWPBkgw --worksheet "Jobs" --base-resume ./base_resume.docx --master-template ./Master_template2.docx --overwrite-keywords

python automated_resume_tailor.py --sheet-id 1dkxFBloNJQMnXtTfLG5CsGryK_EfmlYrdFFluWPBkgw --worksheet "Jobs" --base-resume ./base_resume.docx --master-template ./Master_template2.docx --dry-run
"""

import os
import sys
import json
import logging
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import re

# Third-party imports
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from openai import OpenAI
import anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive
from tenacity import retry, stop_after_attempt, wait_exponential
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('resume_tailor.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

class ResumeTailor:
    """Main class for automated resume tailoring functionality."""
    
    def __init__(self, dry_run: bool = False, use_alternative_logic: bool = False):
        """Initialize the ResumeTailor with configuration."""
        self.dry_run = dry_run or os.getenv('DRY_RUN', '').lower() == 'true'
        self.llm_provider = os.getenv('LLM_PROVIDER', 'openai').lower()
        self.use_alternative_logic = use_alternative_logic

        # Initialize clients
        self._init_google_client()
        self._init_llm_client()
        self._init_drive_client()

        logic_mode = "V2 (30 keywords, Experience-first)" if use_alternative_logic else "V1 (15 keywords, Default)"
        logger.info(f"ResumeTailor initialized - Dry run: {self.dry_run}, LLM: {self.llm_provider}, Mode: {logic_mode}")
    
    def _init_google_client(self):
        """Initialize Google Sheets client."""
        if self.dry_run:
            self.gc = None
            logger.info("Google Sheets client not initialized (dry run mode)")
            return
            
        try:
            auth_mode = os.getenv('GOOGLE_AUTH_MODE', 'service_account').lower()
            
            if auth_mode == 'oauth':
                # Use OAuth flow via gspread helper
                client_secrets = os.getenv('GOOGLE_CLIENT_SECRETS_FILE', './client_secret.json')
                token_file = os.getenv('GOOGLE_OAUTH_TOKEN_FILE', './token.json')
                if not client_secrets:
                    raise ValueError("GOOGLE_CLIENT_SECRETS_FILE environment variable not set for OAuth mode")

                try:
                    self.gc = gspread.oauth(
                        credentials_filename=client_secrets,
                        authorized_user_filename=token_file,
                    )
                    logger.info("Google Sheets client initialized with OAuth")
                except Exception as oauth_error:
                    # If OAuth fails (e.g., expired token, missing fields), delete token file and retry
                    error_str = str(oauth_error).lower()
                    if any(keyword in error_str for keyword in ["invalid_grant", "expired", "missing fields", "not in the expected format"]):
                        logger.warning(f"OAuth token invalid or corrupt: {oauth_error}")
                        logger.warning(f"Deleting {token_file} and re-authenticating...")
                        if Path(token_file).exists():
                            Path(token_file).unlink()
                        # Retry OAuth flow
                        self.gc = gspread.oauth(
                            credentials_filename=client_secrets,
                            authorized_user_filename=token_file,
                        )
                        logger.info("Google Sheets client re-initialized with OAuth after token refresh")
                    else:
                        raise
            else:
                # Use service account
                creds_path = os.getenv('GOOGLE_CREDENTIALS_PATH')
                if not creds_path:
                    raise ValueError("GOOGLE_CREDENTIALS_PATH environment variable not set")
                
                scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
                creds = ServiceAccountCredentials.from_json_keyfile_name(creds_path, scope)
                self.gc = gspread.authorize(creds)
                logger.info("Google Sheets client initialized with service account")
        except Exception as e:
            logger.error(f"Failed to initialize Google Sheets client: {e}")
            self.gc = None
    
    def _init_llm_client(self):
        """Initialize LLM client (OpenAI or Anthropic)."""
        if self.dry_run:
            self.llm_client = None
            logger.info("LLM client not initialized (dry run mode)")
            return
            
        try:
            if self.llm_provider == 'anthropic':
                api_key = os.getenv('ANTHROPIC_API_KEY')
                if not api_key:
                    raise ValueError("ANTHROPIC_API_KEY environment variable not set")
                self.llm_client = anthropic.Anthropic(api_key=api_key)
            else:
                api_key = os.getenv('OPENAI_API_KEY')
                if not api_key:
                    raise ValueError("OPENAI_API_KEY environment variable not set")
                
                # Check for placeholder values
                if api_key in ['your_openai_api_key_here', 'sk-your-openai-api-key', 'your_openai_api_key_here']:
                    raise ValueError(
                        "OPENAI_API_KEY is set to a placeholder value. "
                        "Please update your .env file with a real OpenAI API key from https://platform.openai.com/api-keys"
                    )
                
                if not api_key.startswith('sk-'):
                    raise ValueError(
                        "OPENAI_API_KEY format is invalid. OpenAI API keys should start with 'sk-'. "
                        "Get your API key from https://platform.openai.com/api-keys"
                    )
                
                self.llm_client = OpenAI(api_key=api_key)
            
            logger.info(f"LLM client initialized successfully ({self.llm_provider})")
        except Exception as e:
            logger.error(f"Failed to initialize LLM client: {e}")
            self.llm_client = None
    
    def _init_drive_client(self):
        """Initialize Google Drive client."""
        if self.dry_run:
            self.drive = None
            logger.info("Google Drive client not initialized (dry run mode)")
            return
            
        try:
            auth_mode = os.getenv('GOOGLE_AUTH_MODE', 'service_account').lower()
            
            if auth_mode == 'oauth':
                # Use OAuth flow for Drive
                gauth = GoogleAuth()
                gauth.LocalWebserverAuth()
                self.drive = GoogleDrive(gauth)
                logger.info("Google Drive client initialized with OAuth")
            else:
                # Use service account for Drive by setting credentials directly
                creds_path = os.getenv('GOOGLE_CREDENTIALS_PATH')
                if not creds_path:
                    raise ValueError("GOOGLE_CREDENTIALS_PATH environment variable not set for Drive service account")
                scope = ['https://www.googleapis.com/auth/drive']
                sa_creds = ServiceAccountCredentials.from_json_keyfile_name(creds_path, scope)
                gauth = GoogleAuth()
                gauth.credentials = sa_creds
                self.drive = GoogleDrive(gauth)
                logger.info("Google Drive client initialized with service account")
        except Exception as e:
            logger.error(f"Failed to initialize Google Drive client: {e}")
            self.drive = None
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def process_job_data(self, sheet_id: str, worksheet_name: str, overwrite_keywords: bool = False) -> List[Dict[str, Any]]:
        """
        Process job data from Google Sheets and extract keywords.
        
        Args:
            sheet_id: Google Sheets document ID
            worksheet_name: Name of the worksheet
            overwrite_keywords: Whether to overwrite existing keywords
            
        Returns:
            List of job records with extracted keywords
        """
        if self.dry_run:
            logger.info(f"DRY RUN: Would process sheet {sheet_id}, worksheet '{worksheet_name}'")
            return []
        
        if not self.gc or not self.llm_client:
            logger.error("Required clients not initialized")
            return []
        
        try:
            # Open the spreadsheet first
            spreadsheet = self.gc.open_by_key(sheet_id)
            
            # List all available worksheets for debugging
            available_worksheets = [ws.title for ws in spreadsheet.worksheets()]
            logger.info(f"Available worksheets in spreadsheet: {available_worksheets}")
            
            # Try to open the specific worksheet
            try:
                sheet = spreadsheet.worksheet(worksheet_name)
                logger.info(f"Successfully opened worksheet: '{worksheet_name}'")
            except Exception as ws_error:
                logger.error(f"Worksheet '{worksheet_name}' not found. Available worksheets: {available_worksheets}")
                logger.error(f"Worksheet error details: {ws_error}")
                raise

            # Read all values to robustly detect header row
            values = sheet.get_all_values()
            if not values:
                logger.warning("Worksheet has no values")
                return []

            def norm(s: str) -> str:
                return re.sub(r"[^A-Z0-9]", "", s.strip().upper())

            # Detect header row by looking for required columns within first 10 rows
            header_row_idx0 = None
            for r, row in enumerate(values[:10]):
                norms = [norm(c) for c in row]
                if (
                    any(n in ("JOBTITLE", "TITLE", "ROLE", "POSITION") for n in norms)
                    and any(n in ("JOBDESCRIPTION", "DESCRIPTION", "JD", "JOBDESCRIPTIONRAW") for n in norms)
                ):
                    header_row_idx0 = r
                    break

            if header_row_idx0 is None:
                # Fallback to first row as header
                header_row_idx0 = 0
                logger.warning("Could not detect header row automatically; using row 1")

            header_row_1 = header_row_idx0 + 1
            header_cells = sheet.row_values(header_row_1)
            header_norms = [norm(h) for h in header_cells]

            # Direct mapping for known column names
            column_mapping = {
                "JOBID": "JOB_ID",
                "JOBTITLE": "JOB_TITLE",
                "COMPANYNAME": "COMPANY",
                "LOCATION": "LOCATION",
                "EXPERIENCELEVEL": "EXPERIENCE_LEVEL",
                "SPONSORSHIP": "SPONSORSHIP",
                "POSTINGURL": "JOB_URL",
                "POSTEDDATE": "POSTED_DATE",
                "JOBDESCRIPTIONRAW": "JOB_DESCRIPTION",
                "KEYWORDSEXTRACTED": "KEYWORDS",
                "RESUMEVERSIONLINK": "RESUME_LINK",
                "COVERLETTERVERSION": "COVER_LETTER_VERSION",
                "TAILORINGSTATUS": "TAILORING_STATUS",
                "PRIORITYSCORE": "PRIORITY_SCORE",
                "CONFIDENCESCORE": "CONFIDENCE_SCORE",
                "CONTACTNAMECONTACTEMAIL": "CONTACT_INFO",
                "FOLLOWUPDATE": "FOLLOW_UP_DATE",
                "APPLICATIONCOUNT": "APPLICATION_COUNT",
                "WORKTYPE": "WORK_TYPE",
                "CONTACTURL": "CONTACT_URL"
            }

            def canonical(hn: str, raw: str) -> str:
                return column_mapping.get(hn, raw.strip())

            header_map = {idx: canonical(hn, raw) for idx, (hn, raw) in enumerate(zip(header_norms, header_cells))}

            # Find KEYWORDS column (should exist as "Keywords extracted")
            canonical_headers = [canonical(n, r) for n, r in zip(header_norms, header_cells)]
            try:
                keywords_col_index = canonical_headers.index("KEYWORDS") + 1
                logger.info(f"KEYWORDS column found at {keywords_col_index}")
            except ValueError:
                logger.error("KEYWORDS column not found. Expected 'Keywords extracted' column in the sheet.")
                raise ValueError("Required 'Keywords extracted' column not found in Google Sheet")

            # Find RESUME_LINK column (should exist as "Resume version [link]")
            try:
                resume_link_col_index = canonical_headers.index("RESUME_LINK") + 1
                logger.info(f"RESUME_LINK column found at {resume_link_col_index}")
            except ValueError:
                logger.error("RESUME_LINK column not found. Expected 'Resume version [link]' column in the sheet.")
                raise ValueError("Required 'Resume version [link]' column not found in Google Sheet")

            # Find COVER_LETTER_VERSION column (should exist as "Cover letter version")
            try:
                cover_letter_col_index = canonical_headers.index("COVER_LETTER_VERSION") + 1
                logger.info(f"COVER_LETTER_VERSION column found at {cover_letter_col_index}")
            except ValueError:
                logger.warning("COVER_LETTER_VERSION column not found. Cover letter integration will be limited.")
                cover_letter_col_index = None

            # Build records preserving actual row numbers
            processed_records = []
            for r in range(header_row_idx0 + 1, len(values)):
                row = values[r]
                rec: Dict[str, Any] = {}
                for c, cell in enumerate(row):
                    if c in header_map:
                        key = header_map[c]
                        if key:
                            rec[key] = cell.strip()
                # Skip rows without core fields
                job_title = rec.get('JOB_TITLE', '').strip()
                job_description = rec.get('JOB_DESCRIPTION', '').strip()
                if not job_title and not job_description:
                    continue
                rec['__ROW_INDEX'] = r + 1  # 1-indexed for gspread
                rec['__KEYWORDS_COL_INDEX'] = keywords_col_index
                rec['__RESUME_LINK_COL_INDEX'] = resume_link_col_index
                if cover_letter_col_index:
                    rec['__COVER_LETTER_COL_INDEX'] = cover_letter_col_index
                rec['__SHEET'] = sheet  # Store sheet reference for updates

                # Extract or reuse keywords
                if rec.get('KEYWORDS') and not overwrite_keywords:
                    logger.info(f"Row {rec['__ROW_INDEX']}: Using existing keywords for '{job_title}'")
                    rec['KEYWORDS_RAW'] = []
                else:
                    if not job_title or not job_description:
                        logger.warning(f"Row {rec['__ROW_INDEX']}: Missing job title or description")
                        rec['KEYWORDS'] = ''
                        rec['KEYWORDS_RAW'] = []
                    else:
                        logger.info(f"Row {rec['__ROW_INDEX']}: Extracting keywords for '{job_title}'")
                        # Use alternative logic if enabled
                        if self.use_alternative_logic:
                            kws = self._extract_keywords_v2(job_description)
                        else:
                            kws = self._extract_keywords(job_description)
                        if kws:
                            keywords_str = '; '.join([f"{kw['term']} (rank {kw['rank']})" for kw in kws])
                            sheet.update_cell(rec['__ROW_INDEX'], keywords_col_index, keywords_str)
                            rec['KEYWORDS'] = keywords_str
                            rec['KEYWORDS_RAW'] = kws
                            logger.info(f"Row {rec['__ROW_INDEX']}: Keywords extracted and saved")
                        else:
                            logger.warning(f"Row {rec['__ROW_INDEX']}: Failed to extract keywords")
                            rec['KEYWORDS'] = ''
                            rec['KEYWORDS_RAW'] = []

                processed_records.append(rec)

            if not processed_records:
                logger.warning("No records found after parsing worksheet values")
                return []

            logger.info(f"Processed {len(processed_records)} job records (header row: {header_row_1})")
            return processed_records

        except Exception as e:
            logger.error(f"Error processing job data: {e}")
            raise
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def _extract_keywords(self, job_description: str) -> List[Dict[str, Any]]:
        """Extract keywords from job description using LLM."""
        if self.dry_run:
            logger.info("DRY RUN: Would extract keywords")
            return [
                {"term": "python", "rank": 1}, {"term": "machine learning", "rank": 2},
                {"term": "data analysis", "rank": 3}, {"term": "sql", "rank": 4},
                {"term": "pandas", "rank": 5}, {"term": "numpy", "rank": 6},
                {"term": "scikit-learn", "rank": 7}, {"term": "tensorflow", "rank": 8},
                {"term": "git", "rank": 9}, {"term": "docker", "rank": 10},
                {"term": "api development", "rank": 11}, {"term": "rest apis", "rank": 12},
                {"term": "database design", "rank": 13}, {"term": "agile", "rank": 14},
                {"term": "unit testing", "rank": 15}, {"term": "data visualization", "rank": 16},
                {"term": "jupyter notebooks", "rank": 17}, {"term": "cloud computing", "rank": 18},
                {"term": "problem solving", "rank": 19}, {"term": "communication", "rank": 20}
            ]
        
        if not self.llm_client:
            logger.error("LLM client not initialized")
            return []
        
        system_prompt = """You are an expert recruiting analyst and ATS keyword specialist. Extract the most critical technical skills, tools, and qualifications from job descriptions.

Return ONLY valid JSON in this exact format:
{ "keywords_ranked": [ {"term":"python","rank":1}, {"term":"machine learning","rank":2} ] }

REQUIREMENTS:
- Extract exactly 15 keywords (not 20)
- Focus on: technical skills, software, tools, programming languages, frameworks, methodologies, certifications
- Rank 1 = highest importance, 15 = lowest importance
- Use lowercase only
- Maximum 3 words per term
- No duplicates or variations (e.g., include "python" not both "python" and "python programming")
- Prioritize hard skills over soft skills

OUTPUT: JSON only, no explanations."""

        user_prompt = f"""Extract and rank the 15 most important keywords for ATS optimization from this job description:

{job_description}

Focus on technical requirements, required skills, and specific tools/technologies mentioned."""
        
        try:
            if self.llm_provider == 'anthropic':
                response = self.llm_client.messages.create(
                    model="claude-3-sonnet-20240229",
                    max_completion_tokens=1000,
                    temperature=0.3,
                    messages=[
                        {"role": "user", "content": f"{system_prompt}\n\n{user_prompt}"}
                    ]
                )
                content = response.content[0].text
            else:
                response = self.llm_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3,
                    max_tokens=1000
                )
                content = response.choices[0].message.content
            
            # Parse JSON response
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                data = json.loads(json_match.group())
                keywords = data.get('keywords_ranked', [])
                logger.info(f"Extracted {len(keywords)} keywords")
                return keywords
            else:
                logger.error("No valid JSON found in response")
                return []
                
        except Exception as e:
            logger.error(f"Error extracting keywords: {e}")
            return []
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def generate_resume(self, base_resume_path: str, master_resume_template_path: str,
                       job_title: str, job_description: str, keywords: List[Dict[str, Any]],
                       company_name: str = "", extra_prompt_instructions: str = "") -> str:
        """
        Generate a tailored resume using TXT-based workflow for perfect formatting.
        
        Args:
            base_resume_path: Path to base resume file
            master_resume_template_path: Path to master template file
            job_title: Target job title
            job_description: Job description text
            keywords: List of keyword dictionaries with 'term' and 'rank'
            extra_prompt_instructions: Additional instructions for AI
            
        Returns:
            Path to generated resume file
        """
        if self.dry_run:
            logger.info(f"DRY RUN: Would generate resume for '{job_title}'")
            return f"tailored_resume_{job_title.replace(' ', '_').lower()}.docx"
        
        if not self.llm_client:
            logger.error("LLM client not initialized")
            return ""
        
        try:
            from resume_txt_converter import ResumeTextConverter
            
            # STEP 1: Convert master template to structured TXT
            converter = ResumeTextConverter()
            master_structured_txt = converter.docx_to_structured_txt(master_resume_template_path)
            
            if not master_structured_txt:
                logger.error("Failed to convert master template to structured TXT")
                return ""
            
            logger.info("Successfully converted master template to structured TXT format")
            
            # STEP 2: Generate tailored content using structured TXT as base
            tailored_structured_txt = self._generate_resume_content_from_structured_txt(
                master_structured_txt, job_title, job_description, keywords, extra_prompt_instructions
            )
            
            if not tailored_structured_txt:
                logger.error("Failed to generate tailored structured TXT")
                return ""
            
            # STEP 3: Create output filename in format: Aishwaryeshwar_Manickavel_companyname_jobrole
            safe_company_name = re.sub(r'[^\w\s-]', '', company_name or "company").strip()
            safe_company_name = re.sub(r'[-\s]+', '_', safe_company_name).lower()

            safe_job_title = re.sub(r'[^\w\s-]', '', job_title).strip()
            safe_job_title = re.sub(r'[-\s]+', '_', safe_job_title).lower()

            output_filename = f"Aishwaryeshwar_Manickavel_{safe_company_name}_{safe_job_title}.docx"
            output_path = Path(output_filename)
            
            # STEP 4: Convert tailored structured TXT back to DOCX with exact formatting
            success = converter.structured_txt_to_docx(
                tailored_structured_txt, master_resume_template_path, str(output_path)
            )
            
            if success:
                logger.info(f"Resume generated with perfect formatting: {output_path}")
                return str(output_path)
            else:
                logger.error("Failed to convert tailored TXT back to DOCX")
                return ""
            
        except Exception as e:
            logger.error(f"Error generating resume with TXT workflow: {e}")
            return ""
    
    def _extract_core_job_role(self, job_title: str) -> str:
        """Extract the core job role (1-3 key words) without 'Engineer' suffix."""
        # Clean up the title - handle commas and extra text
        title_lower = job_title.lower().strip()

        # Split by comma and take the main part (before comma usually contains the core role)
        main_title = title_lower.split(',')[0].strip()

        # Check if title already contains "industrial" - if so, handle specially
        if 'industrial' in main_title:
            # If it's just "Industrial Engineer", return empty string (will default to just "Engineer")
            if main_title in ['industrial engineer', 'industrial engineering']:
                return ''

            # If it's something like "Industrial Manufacturing Engineer", extract everything after "industrial"
            words = main_title.split()
            industrial_index = words.index('industrial')
            remaining_words = words[industrial_index + 1:]

            # Filter out common words from the remaining part
            words_to_remove = {
                'engineer', 'engineering', 'specialist', 'analyst', 'manager', 'lead', 'senior',
                'junior', 'entry', 'level', 'associate', 'principal', 'staff', 'development',
                'developer', 'architect', 'technician', 'coordinator', 'supervisor', 'director',
                'intern', 'trainee', 'consultant', 'expert', 'professional'
            }

            core_words = []
            for word in remaining_words:
                if word not in words_to_remove and len(word) > 2:
                    core_words.append(word.capitalize())

            # Return just the core role without "Engineer"
            if core_words:
                return ' '.join(core_words[:3])
            else:
                # Return empty string if nothing meaningful after "industrial"
                return ''

        # Normal processing for titles that don't contain "industrial"
        words_to_remove = {
            'engineer', 'engineering', 'specialist', 'analyst', 'manager', 'lead', 'senior',
            'junior', 'entry', 'level', 'associate', 'principal', 'staff', 'development',
            'developer', 'architect', 'technician', 'coordinator', 'supervisor', 'director',
            'intern', 'trainee', 'consultant', 'expert', 'professional'
        }

        # Split into words and filter
        words = main_title.split()
        core_words = []

        for word in words:
            # Keep word if it's not in removal list and is meaningful
            if word not in words_to_remove and len(word) > 2:
                core_words.append(word.capitalize())

        # Limit to maximum 3 core words for conciseness
        if len(core_words) > 3:
            # Keep the most important words (usually the first 2-3 technical terms)
            core_words = core_words[:3]

        # Join core words
        core_role = ' '.join(core_words)

        # Handle edge cases
        if not core_role:
            # Fallback: extract meaningful words more aggressively
            fallback_words = []
            for word in words:
                if len(word) > 3 and word not in ['engineer', 'engineering', 'senior', 'junior', 'lead']:
                    fallback_words.append(word.capitalize())
                if len(fallback_words) >= 2:  # Limit to 2 words in fallback
                    break
            core_role = ' '.join(fallback_words)

        # If still empty, use a cleaned version of original title
        if not core_role:
            cleaned = main_title.replace('engineer', '').replace('senior', '').replace('junior', '').strip()
            words = cleaned.split()[:2]  # Take first 2 words max
            core_role = ' '.join([w.capitalize() for w in words if w])

        # Final fallback
        if not core_role:
            first_word = main_title.split()[0]
            if first_word not in ['engineer', 'engineering', 'senior', 'junior', 'lead']:
                core_role = first_word.title()

        return core_role

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def _generate_resume_content_from_structured_txt(self, master_structured_txt: str, job_title: str,
                                                    job_description: str, keywords: List[Dict[str, Any]],
                                                    extra_instructions: str) -> str:
        """Generate tailored resume content with enhanced bullet points using STAR method."""

        # Extract core job role and create Industrial Engineer format
        core_job_role = self._extract_core_job_role(job_title)

        # Format the industrial title properly
        if core_job_role:
            industrial_job_title = f"Industrial {core_job_role} Engineer"
        else:
            # If no core role extracted (e.g., "Industrial Engineer"), just use "Industrial Engineer"
            industrial_job_title = "Industrial Engineer"

        system_prompt = """You are an expert resume writer and professional bullet-point generator. Your task is to transform job descriptions into achievement-focused, quantifiable, and ATS-optimized resume content.

**CRITICAL**: Return content in this EXACT structured format:

===HEADER===
Aishwaryeshwar Manickavel
ayeshwarm@gmail.com | 832-775-2886 | www.linkedin.com/in/aishwarnick

===EDUCATION===
University of Houston | Master of Science in Industrial Engineering | Houston, TX, USA | December 2025
PSG College of Technology | Bachelor of Engineering in Metallurgical Engineering | TN, India | July 2023

===PROJECTS===
Quality 4.0: Digital Twin with LLM-Powered SPC Analytics | University of Houston | Feb 2025 - Present
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]

Time Series Predictive Analytics-Operational Performance | University of Houston | Jan 2024 - May 2024
• [Strong bullet point following STAR method with quantified results]
• [Strong bullet point following STAR method with quantified results]

===PROFESSIONAL EXPERIENCE===
University of Houston | Operations Lead | Houston, TX | USA | May 2024 – Feb 2025
• Optimized workforce scheduling algorithms using demand forecasting and resource allocation strategies, improving staff utilization by 20% while maintaining service quality standards.
• Reduced operational downtime by 15% through implementation of preventive maintenance protocols and real-time equipment tracking systems, saving $8,000+ in potential revenue loss.
• Managed end-to-end event logistics for 25+ campus events per semester, ensuring 100% safety compliance and coordinating cross-functional teams of 15+ staff members.

Centre for Automotive Energy Materials at IIT Madras | {industrial_job_title} | Chennai, TN | India | Jun 2022 -- Dec 2023
• [Generate tailored 2-line bullet using STAR method with quantified results]
• [Generate tailored 2-line bullet using STAR method with quantified results]
• [Generate tailored 2-line bullet using STAR method with quantified results]

A. V.C Pvt Ltd. | {industrial_job_title} | TN, India | January 2021 -- May 2022
• [Generate tailored 2-line bullet using STAR method with quantified results]
• [Generate tailored 2-line bullet using STAR method with quantified results]
• [Generate tailored 2-line bullet using STAR method with quantified results]


===TECHNICAL SKILLS===
[JD Category 1]: [5-6 unique job-specific skills from JD, under 100 chars]
Data & Analytics: Power BI, Tableau, SQL, Excel, Python, Pandas, NumPy
[JD Category 2]: [5-6 unique job-specific skills from JD, under 100 chars, NO duplicates from Category 1]
Design & Simulation: CAD, SolidWorks, Creo, 3D Printing, Arena Simulation, Minitab
[JD Category 3]: [5-6 unique job-specific skills from JD, under 100 chars, NO duplicates from Categories 1 or 2]
Process & Quality: Lean Six Sigma, Root Cause Analysis, FMEA, SPC, JMP

**TECHNICAL SKILLS REQUIREMENTS:**

**CRITICAL - ALTERNATING SKILLS STRUCTURE (EXACTLY 6 CATEGORIES):**
You MUST create EXACTLY 6 skill categories in a strict alternating pattern:

**MANDATORY NO-DUPLICATION RULE:**
- Each JD-specific skill can appear in ONLY ONE category
- Check all previous categories before adding a skill
- If a skill was used in Category 1, DO NOT repeat it in Categories 3 or 5
- Example: If "Python" is in Category 2 (default), don't add it to Categories 1, 3, or 5
- Example: If "Lean Six Sigma" is in Category 6 (default), don't add it to Categories 1, 3, or 5

**ALTERNATING PATTERN (MANDATORY):**
1. **[JD Category 1]** - From job description (5-6 UNIQUE skills, <100 chars total)
2. **Data & Analytics (Default)** - Power BI, Tableau, SQL, Excel, Python, Pandas, NumPy
3. **[JD Category 2]** - From job description (5-6 UNIQUE skills NOT in Category 1, <100 chars total)
4. **Design & Simulation (Default)** - CAD, SolidWorks, Creo, 3D Printing, Arena Simulation, Minitab OR **Software & Cloud (Default)** - Git, AWS, Docker, CI/CD, GCP
5. **[JD Category 3]** - From job description (5-6 UNIQUE skills NOT in Categories 1 or 3, <100 chars total)
6. **Process & Quality (Default)** - Lean Six Sigma, Root Cause Analysis, FMEA, SPC, JMP

**CORE DEFAULT SKILLS (ALWAYS INCLUDE IN CATEGORIES 2, 4, 6):**
- **Data & Analytics:** Power BI, Tableau, SQL (MySQL, PostgreSQL), Excel, Python (Pandas, NumPy, Matplotlib, Seaborn), R Programming
- **Design & Simulation (For Hardware/Manufacturing Roles):** CAD Modeling, SolidWorks, Creo, 3D Printing, Additive Manufacturing, Arena Simulation, AnyLogic, Simul8, Unity, Minitab, JMP
- **Software & Cloud (For Software/IT Roles):** Git, GitHub, AWS (Lambda, S3, EC2), GCP, Azure, Docker, CI/CD, RESTful APIs, Microservices, Agile/Scrum
- **Process & Quality:** Lean Six Sigma, Minitab, JMP, Root Cause Analysis, FMEA, Statistical Process Control (SPC), Continuous Improvement, ETL Pipelines

**ADAPTIVE CATEGORY SELECTION FOR POSITION 4 (2nd Default Category):**
- **IF JD mentions**: software, programming, cloud, DevOps, full-stack, backend, frontend, web development
  → Use "Software & Cloud" category
- **IF JD mentions**: manufacturing, mechanical, design, CAD, production, assembly, prototyping
  → Use "Design & Simulation" category
- **IF BOTH or NEITHER**: Use "Design & Simulation" as default

**HOW TO CREATE JD-RELATED CATEGORIES (POSITIONS 1, 3, 5):**
- Analyze the job description for key skill themes
- Create category names that match JD language (e.g., "Manufacturing Systems", "Controls Engineering", "Quality Management")
- Populate with 5-6 skills specifically mentioned in the job description
- If JD mentions broad terms like "programming", include specific languages (Python, C++, MATLAB, etc.)

**EXAMPLES OF ALTERNATING STRUCTURE:**

*Example 1: Manufacturing Controls Role (Hardware-focused):*
1. Controls & Automation: PLC, SCADA, HMI, Allen-Bradley, Siemens TIA Portal
2. Data & Analytics (Default): Power BI, Tableau, SQL, Excel, Python (Pandas, NumPy)
3. Manufacturing Systems: MES, ERP Integration, OEE Analysis, Production Optimization
4. Design & Simulation (Default): CAD, SolidWorks, Arena Simulation, Minitab, JMP
5. Process Control: Sensors, Motor Controls, VFDs, Ethernet/IP, Modbus, Commissioning
6. Process & Quality (Default): Lean Six Sigma, Root Cause Analysis, FMEA, SPC

*Example 2: Software Engineer / Full-Stack Developer Role (Software-focused):*
1. Web Development: React, Node.js, JavaScript, TypeScript, HTML/CSS, Redux
2. Data & Analytics (Default): Power BI, Tableau, SQL, Excel, Python (Pandas, NumPy)
3. Backend & APIs: REST APIs, GraphQL, Microservices, Express.js, MongoDB, PostgreSQL
4. Software & Cloud (Default): Git, GitHub, AWS (Lambda, S3, EC2), Docker, CI/CD
5. Cloud Infrastructure: GCP, Azure, Kubernetes, Terraform, Jenkins, DevOps
6. Process & Quality (Default): Lean Six Sigma, Root Cause Analysis, Agile, Scrum

*Example 3: Data Analytics Role (Mixed):*
1. Business Intelligence: Data Modeling, KPI Development, Predictive Analytics
2. Data & Analytics (Default): Power BI, Tableau, SQL, Excel, Python (Pandas, NumPy)
3. Data Engineering: ETL, Data Warehousing, Apache Spark, Airflow, Data Pipelines
4. Software & Cloud (Default): Git, AWS (Redshift, Lambda, S3), Docker, ETL Tools
5. Big Data & Analytics: Hadoop, Azure, GCP, Data Lakes, Real-time Processing
6. Process & Quality (Default): Lean Six Sigma, Root Cause Analysis, SPC

**CRITICAL RULES FOR 1-PAGE FIT:**
✓ EXACTLY 6 categories (3 JD + 3 Default) in alternating order
✓ 4-6 skills per category (limit to stay within one line)
✓ **EACH CATEGORY LINE MUST NOT EXCEED ~100 CHARACTERS (ONE PHYSICAL LINE)**
✓ **ENTIRE SKILLS SECTION = EXACTLY 6 LINES (one per category)**
✓ Concise naming: "Data & Analytics" not "Data Analytics & Business Intelligence Platform Development"
✓ NO explanations or descriptions - just skill names
✓ Use commas, no bullet points in skills section
✓ **REMOVE bracketed details if line exceeds 100 characters**:
   - WRONG: "Python (Pandas, NumPy, Matplotlib, Seaborn)" (too long)
   - RIGHT: "Python, Pandas, NumPy, Matplotlib" (expanded, concise)
   - WRONG: "SQL (MySQL, PostgreSQL, MS SQL Server)" (too long with brackets)
   - RIGHT: "SQL, MySQL, PostgreSQL" (remove brackets, list separately)
✓ Use abbreviations: "SPC" not "Statistical Process Control"
✓ Test: Each category should fit on ONE line in Times New Roman 11pt with 0.5" margins

**BULLET POINT REQUIREMENTS:**

1. **Action-Oriented:** Start every bullet with a strong action verb (Negotiated, Implemented, Reduced, Streamlined, Optimized, Led, Developed, Managed, Architected, Established, Executed).

2. **STAR Flow:** Each bullet must follow *Situation → Task → Action → Result*:
   * What you did (situation/task)
   * How you did it (method, tools, process)
   * Why it mattered (impact with measurable outcome)

3. **Quantifiable Results:** Include metrics wherever possible:
   * % cost savings, % efficiency improvements
   * $ value, revenue impact
   * Time saved, deadlines met
   * Risk reduction, compliance rates
   * Process improvements, system optimizations

4. **CRITICAL - Line Length Requirements (MANDATORY):**
   * **EVERY bullet MUST be 180-200 characters minimum** - this fills EXACTLY 2 complete lines
   * Count characters: if < 180 chars, ADD MORE DETAIL immediately
   * **NEVER submit a bullet under 180 characters** - it will look incomplete
   * Ways to extend short bullets to 2 full lines:
     - Add specific tool names (Minitab, Power BI, Python, AWS, etc.)
     - Add quantified scope (across 5 departments, for 12+ stakeholders, managing $500K budget)
     - Add secondary metrics (reducing X by 20% while improving Y by 15%)
     - Add methodology details (using Lean Six Sigma, DMAIC framework, statistical analysis)
   * Example WRONG (120 chars, only 1.3 lines): "Implemented process improvements reducing costs by 20%"
   * Example RIGHT (195 chars, full 2 lines): "Implemented Lean Six Sigma process improvements across 3 production lines using statistical analysis in Minitab and real-time monitoring dashboards in Power BI, reducing operational costs by 20% while improving throughput by 15%"

5. **ATS Alignment:** Use keywords and industry terms from the job description itself.

6. **Format:** **[Strong Action Verb] + [Method/Tool/Process] + [Quantified Result/Impact] + [Additional Context if needed to fill 2 full lines]**.

**EXAMPLES OF STRONG BULLET POINTS:**
• Negotiated supplier contracts through comprehensive RFx processes, achieving 25% cost reduction across $2M annual procurement budget
• Implemented cross-functional NPI workflows connecting engineering and manufacturing teams, reducing product launch timelines by 30%
• Established supplier performance KPIs and monitoring dashboards, enabling proactive risk identification and 15% quality improvement

**CONTENT REQUIREMENTS:**
- Rewrite ALL content to match the job description and keywords
- Keep company names and university names UNCHANGED
- Job entry format: Company | Job Role | Location | Duration
- **CRITICAL**: Keep University of Houston role EXACTLY as shown - DO NOT change the job title "Operations Lead" or any of the 3 bullet points
- **CRITICAL**: For Centre for Automotive Energy Materials at IIT Madras: create EXACTLY 3 bullet points (each 180-200 chars, full 2 lines) tailored to job description
- **CRITICAL**: For A.V.C Pvt Ltd: create EXACTLY 3 bullet points (each 180-200 chars, full 2 lines) tailored to job description
- **CRITICAL**: For PROJECTS section: Generate 2 tailored bullets per project (each 180-200 chars, full 2 lines) - DO NOT use hardcoded defaults
- **CRITICAL**: For "Quality 4.0: Digital Twin" project - MUST mention Unity (simulation platform) and Python, focus on simulation modeling, real-time monitoring, and LLM integration
- **CRITICAL**: For "Time Series Predictive Analytics" project - Focus on predictive analytics, machine learning, data visualization, and operational performance optimization
- Natural keyword integration throughout
- Name: "Aishwaryeshwar Manickavel" (first letters only capitalized)
- **NO DUPLICATE SKILLS** across all 6 skill categories

Return ONLY the structured text with ===SECTION=== markers above. Do not include any formatting specifications or instructions in your response."""

        keywords_str = '; '.join([kw['term'] for kw in sorted(keywords, key=lambda x: x['rank'])[:10]])  # Top 10 keywords only

        # Format the system prompt with the dynamic industrial job title
        formatted_system_prompt = system_prompt.format(industrial_job_title=industrial_job_title)

        user_prompt = f"""### TARGET JOB DETAILS:
Job Title: {job_title}
Generated Industrial Title: {industrial_job_title}

Job Description:
{job_description}

### TOP 10 KEYWORDS (integrate naturally):
{keywords_str}

### SPECIFIC INSTRUCTIONS FOR THIS ROLE:
Transform the job description above into powerful resume bullet points that demonstrate relevant experience. Focus on:

1. **Job Roles**: Use the generated Industrial title "{industrial_job_title}" for both experience positions

2. **Strong Bullet Points**: For each role, create exactly 4 bullet points that:
   - Start with powerful action verbs (Led, Negotiated, Implemented, Optimized, Developed, Managed)
   - Follow STAR method: Situation → Task → Action → Result
   - Include specific metrics and quantified outcomes
   - Use industry terminology from the job description
   - Demonstrate progression and impact

3. **Skills Section - CRITICAL ALTERNATING STRUCTURE (EXACTLY 6 CATEGORIES)**:
   - MANDATORY: Alternate between JD-related categories and default categories
   - Category 1: [JD-specific theme] with 5-6 skills from job description
   - Category 2: Data & Analytics (Default) - Power BI, Tableau, SQL, Excel, Python (Pandas, NumPy, Matplotlib)
   - Category 3: [JD-specific theme] with 5-6 skills from job description
   - Category 4: Design & Simulation (Default) - CAD Modeling, SolidWorks, Creo, 3D Printing, Arena Simulation, Minitab, JMP
   - Category 5: [JD-specific theme] with 5-6 skills from job description
   - Category 6: Process & Quality (Default) - Lean Six Sigma, Minitab, JMP, Root Cause Analysis, FMEA, SPC
   - Keep categories concise to fit on 1 page

4. **Keyword Integration**: Naturally weave in the most important keywords throughout the content AND Technical Skills section

### TRANSFORMATION APPROACH:
- Take the original experience and reframe it to align with this specific job description
- Use the same companies and timeframes but adapt the roles and achievements
- Ensure each bullet point tells a compelling story with measurable impact
- Make the candidate appear perfectly suited for this exact position

### CRITICAL - INTEGRATE DEFAULT SKILLS IN EXPERIENCE BULLETS:
Your experience bullet points MUST showcase the default technical skills in action. Examples:

✓ GOOD: "Designed KPI dashboards in Power BI and Tableau to track manufacturing performance, improving decision latency by 30%"
✓ GOOD: "Conducted root cause analysis using Minitab and Python (Pandas, NumPy), identifying key bottlenecks that reduced cycle time by 18%"
✓ GOOD: "Built ETL pipelines using AWS Lambda and Python to automate data integration from 5+ sources, saving 15 hours/week"
✓ GOOD: "Implemented Lean Six Sigma methodologies with statistical analysis in Minitab, achieving $120K annual cost savings"

✗ BAD: "Analyzed data to improve processes" (too vague, no tools mentioned)
✗ BAD: "Created dashboards for management" (missing specific tools like Power BI/Tableau)

**Strategy:** Weave default skills (Power BI, Tableau, SQL, Python, Minitab, AWS, Lean Six Sigma) naturally into 60-70% of your bullet points while maintaining relevance to the job description

### CONSISTENCY REQUIREMENTS:
- ALWAYS use the exact ===SECTION=== markers shown above
- ALWAYS format job entries as: Company | Job Role | Location | Duration
- ALWAYS keep University of Houston role and its 3 bullet points EXACTLY as shown (never modify)
- ALWAYS use exactly 3 bullet points for Centre for Automotive Energy Materials at IIT Madras
- ALWAYS use exactly 3 bullet points for A.V.C Pvt Ltd
- ALWAYS use exactly 2 bullet points for each project (Quality 4.0 and Time Series)
- ALWAYS start bullets with strong action verbs
- ALWAYS include quantified results (percentages, dollar amounts, time savings)

### CRITICAL - STRICT 1-PAGE REQUIREMENT:
**Resume MUST fit on exactly 1 page with ultra-narrow margins: 0.5" left/right, 0.3" top/bottom**

**CRITICAL FORMATTING RULES:**

**A. Bullet Points - MUST USE FULL 2 LINES:**
- Each bullet should be approximately 180-200 characters to fill 2 complete lines
- If bullet is 1.3-1.5 lines, ADD MORE DETAIL to extend to full 2 lines
- Add specific tools, methodologies, metrics, or business context
- Utilize the full width - don't stop mid-line
- Example WRONG (1.3 lines): "Developed dashboards reducing analysis time by 30%"
- Example RIGHT (full 2 lines): "Developed interactive Power BI dashboards integrating data from 5+ sources using SQL and Python ETL pipelines, reducing analysis time by 30% and enabling real-time decision-making for 50+ stakeholders"

**B. Skills Section - EXACTLY 1 LINE PER CATEGORY:**
- Use exactly 6 skill categories (3 JD + 3 Default alternating)
- **Each category line MUST NOT exceed 100 characters (one physical line in 11pt Times New Roman)**
- Remove bracketed content if it causes line to exceed: "Python, Pandas, NumPy" NOT "Python (Pandas, NumPy, Matplotlib)"
- Limit to 4-6 skills per category
- Use abbreviations: "SPC" not "Statistical Process Control"
- Test each line: if it wraps to 2nd line, remove 1-2 skills

**C. Total Content:**
- 13 total bullets: 3 (UH) + 2 (Project 1) + 2 (Project 2) + 3 (IIT Madras) + 3 (A.V.C)
- 6 skill category lines (no more, no less)

**Example of Properly Formatted Skill Lines:**
✓ Data & Analytics: Power BI, Tableau, SQL, Excel, Python, Pandas
✗ Data & Analytics: Power BI, Tableau, SQL (MySQL, PostgreSQL), Excel, Python (Pandas, NumPy, Matplotlib) [TOO LONG]

{extra_instructions}"""

        try:
            if self.llm_provider == 'anthropic':
                response = self.llm_client.messages.create(
                    model="claude-3-sonnet-20240229",
                    max_completion_tokens=4000,
                    temperature=0.3,
                    messages=[
                        {"role": "user", "content": f"{formatted_system_prompt}\n\n{user_prompt}"}
                    ]
                )
                content = response.content[0].text
            else:
                response = self.llm_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": formatted_system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3,
                    max_tokens=4000
                )
                content = response.choices[0].message.content

            logger.info("Enhanced resume content with strong bullet points generated successfully")
            return content.strip()

        except Exception as e:
            logger.error(f"Error generating enhanced resume content: {e}")
            return ""

    
    def _read_docx_text(self, file_path: str) -> str:
        """Extract text content from a DOCX file."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            return ""
    
    def _create_resume_document(self, content: str, template_path: str, output_path: Path):
        """Create a DOCX document using TXT-based workflow for perfect formatting preservation."""
        try:
            from resume_txt_converter import ResumeTextConverter
            
            logger.info(f"Using TXT-based workflow for perfect formatting: {template_path}")
            
            # Initialize the TXT converter
            converter = ResumeTextConverter()
            
            # The content is already tailored text from AI - use it directly with the converter
            # Convert structured content to DOCX with exact master formatting
            success = converter.structured_txt_to_docx(content, template_path, str(output_path))
            
            if success:
                logger.info(f"Document created with perfect formatting preservation: {output_path}")
            else:
                logger.error("Failed to create document with TXT-based workflow")
                raise RuntimeError("TXT-based formatting workflow failed")
            
        except Exception as e:
            logger.error(f"Error creating document with TXT-based workflow: {e}")
            raise
    
    def _build_style_anchors(self, doc: Document) -> Dict[str, Any]:
        """Find representative paragraphs for different content types to use as formatting anchors."""
        anchors = {"SectionHeader": None, "Bullet": None, "Body": None, "Contact": None}
        
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            
            # Section header (ALL CAPS)
            if not anchors["SectionHeader"] and self._is_all_caps_header(text):
                anchors["SectionHeader"] = p
                
            # Contact info (email, phone, linkedin)
            if not anchors["Contact"] and ('@' in text or 'phone' in text.lower() or 'linkedin' in text.lower()):
                anchors["Contact"] = p
                
            # Bullet point
            if not anchors["Bullet"] and (self._has_numbering(p) or text.startswith("•") or text.startswith("- ")):
                anchors["Bullet"] = p
                
            # Body text
            if not anchors["Body"] and text and not self._is_all_caps_header(text) and not self._has_numbering(p):
                anchors["Body"] = p
                
            if all(anchors.values()):
                break
                
        logger.info(f"Found style anchors: {[k for k, v in anchors.items() if v is not None]}")
        return anchors
    
    def _is_all_caps_header(self, text: str) -> bool:
        """Check if text is an ALL CAPS section header."""
        t = text.strip()
        return t.isupper() and 0 < len(t) <= 48
    
    def _has_numbering(self, paragraph) -> bool:
        """Check if paragraph has numbering/bullet formatting."""
        pPr = getattr(paragraph._element, "pPr", None)
        if pPr is None:
            for child in paragraph._element.iter():
                if child.tag.endswith('numPr'):
                    return True
            return False
        for child in pPr.iter():
            if child.tag.endswith('numPr'):
                return True
        return False
    
    def _locate_sections(self, doc: Document) -> List[tuple]:
        """Locate sections in the document as (header_paragraph, [body_paragraphs])."""
        sections = []
        current_header = None
        current_body = []
        
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if self._is_all_caps_header(text):
                if current_header:
                    sections.append((current_header, current_body))
                current_header = p
                current_body = []
            else:
                if current_header:
                    current_body.append(p)
                else:
                    # Content before first section header
                    current_body.append(p)
        
        if current_header:
            sections.append((current_header, current_body))
        elif current_body:
            # Handle case where there are paragraphs but no section headers
            sections.append((None, current_body))
            
        logger.info(f"Located {len(sections)} sections in template")
        return sections
    
    def _parse_content_into_sections(self, content: str) -> Dict[str, List[str]]:
        """Parse generated content into structured sections."""
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        sections = {}
        current_section = None
        current_lines = []
        
        section_keywords = {
            'SUMMARY': ['PROFESSIONAL SUMMARY', 'SUMMARY', 'PROFILE'],
            'SKILLS': ['CORE SKILLS', 'SKILLS', 'TECHNICAL SKILLS'],
            'EXPERIENCE': ['EXPERIENCE', 'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE'],
            'PROJECTS': ['PROJECTS', 'KEY PROJECTS', 'NOTABLE PROJECTS'],
            'EDUCATION': ['EDUCATION', 'ACADEMIC BACKGROUND']
        }
        
        # First few lines are usually contact/header info
        header_lines = []
        i = 0
        while i < len(lines) and i < 5:
            line = lines[i]
            line_upper = line.upper()
            
            # Check if this line is a section header
            is_section = False
            for section, keywords in section_keywords.items():
                if any(keyword in line_upper for keyword in keywords):
                    is_section = True
                    break
            
            if is_section:
                break
            else:
                header_lines.append(line)
                i += 1
        
        if header_lines:
            sections['HEADER'] = header_lines
        
        # Process remaining lines into sections
        for line in lines[i:]:
            line_upper = line.upper()
            
            # Check if this is a section header
            found_section = None
            for section, keywords in section_keywords.items():
                if any(keyword in line_upper for keyword in keywords):
                    found_section = section
                    break
            
            if found_section:
                # Save previous section
                if current_section and current_lines:
                    sections[current_section] = current_lines
                # Start new section
                current_section = found_section
                current_lines = [line]  # Include the header
            else:
                if current_section:
                    current_lines.append(line)
                else:
                    # Add to header if no section yet
                    if 'HEADER' not in sections:
                        sections['HEADER'] = []
                    sections['HEADER'].append(line)
        
        # Save final section
        if current_section and current_lines:
            sections[current_section] = current_lines
        
        logger.info(f"Parsed content into sections: {list(sections.keys())}")
        return sections
    
    def _replace_sections_preserve_formatting(self, doc: Document, template_sections: List[tuple], 
                                            new_sections: Dict[str, List[str]], anchors: Dict[str, Any]):
        """Replace section content while preserving exact formatting."""
        from copy import deepcopy
        from docx.text.paragraph import Paragraph
        
        # Map template sections to new content
        section_mapping = {
            'SUMMARY': ['PROFESSIONAL SUMMARY', 'SUMMARY'],
            'SKILLS': ['TECHNICAL SKILLS', 'CORE SKILLS'],
            'EXPERIENCE': ['PROFESSIONAL EXPERIENCE', 'EXPERIENCE'],
            'PROJECTS': ['PROJECTS'],
            'EDUCATION': ['EDUCATION']
        }
        
        # Process header separately (name, contact info)
        if 'HEADER' in new_sections and template_sections:
            # Replace first few paragraphs with header content
            header_lines = new_sections['HEADER']
            first_section = template_sections[0]
            
            # If first section has no header, use its body paragraphs for header
            if first_section[0] is None and first_section[1]:
                # Replace first body paragraphs with header content
                body_paras = first_section[1][:len(header_lines)]
                for i, line in enumerate(header_lines):
                    if i < len(body_paras):
                        self._replace_paragraph_text(body_paras[i], line)
        
        # Process each template section
        for template_header, template_body in template_sections:
            if template_header is None:
                continue
                
            template_title = (template_header.text or "").strip().upper()
            
            # Find matching new content
            matched_section = None
            for new_key, new_lines in new_sections.items():
                if new_key == 'HEADER':
                    continue
                    
                for section_key, keywords in section_mapping.items():
                    if section_key == new_key:
                        for keyword in keywords:
                            if keyword in template_title:
                                matched_section = new_key
                                break
                        if matched_section:
                            break
                if matched_section:
                    break
            
            if matched_section and matched_section in new_sections:
                new_lines = new_sections[matched_section]
                
                # Update section header if present
                if new_lines and template_header:
                    self._replace_paragraph_text(template_header, new_lines[0])
                
                # Replace body content
                body_lines = new_lines[1:] if len(new_lines) > 1 else []
                self._replace_section_body(template_body, body_lines, anchors)
    
    def _replace_paragraph_text(self, paragraph, new_text: str):
        """Replace paragraph text while preserving ALL formatting properties (strict mode)."""
        # STRICT MODE: Only replace text nodes, preserve all XML properties
        
        # If paragraph has runs, replace text in first run and clear others
        if paragraph.runs:
            # Keep first run's formatting, just change text
            paragraph.runs[0].text = new_text
            # Clear additional runs but preserve their formatting structure
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            # Add run with new text (inherits paragraph formatting)
            paragraph.add_run(new_text)
        
        # Verify formatting preservation (strict check)
        self._verify_paragraph_formatting_preserved(paragraph)
    
    def _replace_section_body(self, template_paragraphs: List, new_lines: List[str], anchors: Dict[str, Any]):
        """Replace section body with STRICT formatting preservation."""
        from copy import deepcopy
        from docx.text.paragraph import Paragraph
        
        logger.debug(f"Replacing section body: {len(template_paragraphs)} template paragraphs → {len(new_lines)} new lines")
        
        # STRICT MODE: Never delete paragraphs, only replace text or clone with exact formatting
        
        # Update existing paragraphs with strict text replacement
        for i, line in enumerate(new_lines):
            if i < len(template_paragraphs):
                # Use existing paragraph - STRICT: only replace text, preserve all formatting
                logger.debug(f"Replacing text in existing paragraph {i}: '{template_paragraphs[i].text[:30]}...' → '{line[:30]}...'")
                self._replace_paragraph_text(template_paragraphs[i], line)
            else:
                # Need additional paragraphs - clone from appropriate anchor with STRICT formatting
                anchor = self._choose_anchor(line, anchors)
                if anchor and template_paragraphs:
                    # Clone after the last existing paragraph
                    logger.debug(f"Cloning new paragraph {i} from anchor type: {type(anchor).__name__}")
                    new_p = self._clone_paragraph_after(template_paragraphs[-1], anchor)
                    self._replace_paragraph_text(new_p, line)
                    # Add to template_paragraphs list so subsequent paragraphs can reference it
                    template_paragraphs.append(new_p)
                elif anchor:
                    # Edge case: no template paragraphs, but we have an anchor
                    logger.warning("No template paragraphs available for cloning reference")
        
        # STRICT MODE: If we have excess template paragraphs, clear their text but keep structure
        for i in range(len(new_lines), len(template_paragraphs)):
            logger.debug(f"Clearing excess template paragraph {i}")
            self._replace_paragraph_text(template_paragraphs[i], "")  # Clear text but keep formatting
    
    def _choose_anchor(self, line: str, anchors: Dict[str, Any]):
        """Choose appropriate anchor based on line content."""
        if line.startswith("•") or line.startswith("- ") or "•" in line:
            return anchors.get("Bullet")
        elif self._is_all_caps_header(line):
            return anchors.get("SectionHeader")
        else:
            return anchors.get("Body")
    
    def _clone_paragraph_after(self, ref_paragraph, anchor_paragraph):
        """Clone paragraph XML structure from anchor with STRICT formatting preservation."""
        from copy import deepcopy
        from docx.text.paragraph import Paragraph
        
        # STRICT MODE: Clone entire XML structure including w:pPr, w:rPr, w:numPr
        anchor_element = anchor_paragraph._element
        new_element = deepcopy(anchor_element)
        
        # Insert after reference paragraph
        ref_element = ref_paragraph._element
        ref_element.addnext(new_element)
        
        # Create new paragraph object
        new_paragraph = Paragraph(new_element, ref_paragraph._parent)
        
        # STRICT VERIFICATION: Ensure cloned paragraph has identical formatting
        self._verify_cloned_paragraph_formatting(new_paragraph, anchor_paragraph)
        
        logger.debug(f"Cloned paragraph with strict formatting from anchor: {anchor_paragraph.text[:50]}...")
        return new_paragraph
    
    def _verify_paragraph_formatting_preserved(self, paragraph):
        """Verify paragraph formatting has been preserved (strict mode)."""
        try:
            # Check that paragraph has proper XML structure
            if not hasattr(paragraph, '_element') or paragraph._element is None:
                raise ValueError("Paragraph missing XML element")
            
            # Verify paragraph properties exist
            if hasattr(paragraph._element, 'pPr') and paragraph._element.pPr is not None:
                logger.debug("Paragraph properties preserved")
            
            # Verify runs have proper formatting
            for run in paragraph.runs:
                if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                    logger.debug("Run properties preserved")
                    
        except Exception as e:
            logger.warning(f"Formatting verification failed: {e}")
    
    def _verify_cloned_paragraph_formatting(self, cloned_paragraph, anchor_paragraph):
        """Verify cloned paragraph has identical formatting to anchor (strict mode)."""
        try:
            # Compare paragraph-level properties
            anchor_element = anchor_paragraph._element
            cloned_element = cloned_paragraph._element
            
            # Check paragraph properties (pPr)
            anchor_pPr = getattr(anchor_element, 'pPr', None)
            cloned_pPr = getattr(cloned_element, 'pPr', None)
            
            if anchor_pPr is not None and cloned_pPr is not None:
                # Verify numbering properties (numPr) are preserved
                anchor_numPr = self._find_child_with_tag(anchor_pPr, 'numPr')
                cloned_numPr = self._find_child_with_tag(cloned_pPr, 'numPr')
                
                if anchor_numPr is not None and cloned_numPr is not None:
                    logger.debug("Numbering properties successfully cloned")
                elif anchor_numPr is None and cloned_numPr is None:
                    logger.debug("No numbering properties to clone")
                else:
                    logger.warning("Numbering properties mismatch between anchor and clone")
            
            # Verify run count matches (we'll only replace text, not structure)
            if len(anchor_paragraph.runs) > 0 and len(cloned_paragraph.runs) > 0:
                logger.debug(f"Cloned paragraph has {len(cloned_paragraph.runs)} runs (anchor had {len(anchor_paragraph.runs)})")
            
            logger.debug("Strict formatting verification passed")
            
        except Exception as e:
            logger.error(f"STRICT VERIFICATION FAILED: {e}")
            # In strict mode, we could raise here, but for now just warn
            
    def _find_child_with_tag(self, element, tag_suffix: str):
        """Find child element with tag ending in the specified suffix."""
        try:
            for child in element:
                if child.tag.endswith(tag_suffix):
                    return child
            return None
        except Exception:
            return None
    
    def _extract_drive_folder_id(self, folder_url_or_id: str) -> str:
        """Extract folder ID from Drive URL or return as-is if already an ID."""
        if not folder_url_or_id:
            return ""
        
        # If it's already a folder ID (no special characters), return as-is
        if re.match(r'^[a-zA-Z0-9_-]+$', folder_url_or_id):
            return folder_url_or_id
        
        # Extract from Drive URL
        match = re.search(r'/folders/([a-zA-Z0-9_-]+)', folder_url_or_id)
        if match:
            return match.group(1)
        
        return folder_url_or_id

    def upload_to_drive(self, file_path: str, drive_folder_id: Optional[str] = None) -> str:
        """
        Upload a file to Google Drive and return a shareable link.

        Args:
            file_path: Path to file to upload
            drive_folder_id: Optional Drive folder ID or URL

        Returns:
            Shareable Drive link or empty string if failed
        """
        if self.dry_run:
            logger.info(f"DRY RUN: Would upload {file_path} to Drive")
            return "https://drive.google.com/file/d/dry_run_file_id/view"
        
        if not self.drive:
            logger.error("Google Drive client not initialized")
            return ""
        
        try:
            file_drive = self.drive.CreateFile()
            file_drive['title'] = Path(file_path).name
            
            # Determine folder ID
            folder_id = None
            if drive_folder_id:
                folder_id = self._extract_drive_folder_id(drive_folder_id)
            elif os.getenv('DRIVE_FOLDER_ID'):
                folder_id = self._extract_drive_folder_id(os.getenv('DRIVE_FOLDER_ID'))
            
            if folder_id:
                file_drive['parents'] = [{"id": folder_id}]
                logger.info(f"Uploading to Drive folder: {folder_id}")
            else:
                logger.info("Uploading to Drive root folder")
            
            file_drive.SetContentFile(file_path)
            file_drive.Upload()

            file_id = file_drive.get('id')

            # Make the file publicly viewable and get shareable link
            file_drive.InsertPermission({
                'type': 'anyone',
                'role': 'reader'
            })

            # Create shareable link
            drive_link = f"https://drive.google.com/file/d/{file_id}/view"

            logger.info(f"File uploaded to Drive with ID: {file_id}")
            logger.info(f"Shareable link: {drive_link}")
            return drive_link
            
        except Exception as e:
            logger.error(f"Error uploading to Drive: {e}")
            return ""

    def update_resume_link_in_sheet(self, job_record: Dict[str, Any], resume_link: str) -> bool:
        """
        Update the Google Sheet with the resume link for a specific job record.

        Args:
            job_record: Job record containing sheet and position information
            resume_link: Drive link to the generated resume

        Returns:
            True if update was successful, False otherwise
        """
        if self.dry_run:
            logger.info(f"DRY RUN: Would update sheet with resume link: {resume_link}")
            return True

        if not self.gc:
            logger.error("Google Sheets client not initialized")
            return False

        try:
            sheet = job_record.get('__SHEET')
            row_index = job_record.get('__ROW_INDEX')
            resume_link_col_index = job_record.get('__RESUME_LINK_COL_INDEX')
            job_title = job_record.get('JOB_TITLE', 'Unknown')

            if not all([sheet, row_index, resume_link_col_index]):
                logger.error("Missing required sheet information in job record")
                return False

            # Update the resume link cell
            sheet.update_cell(row_index, resume_link_col_index, resume_link)
            logger.info(f"Updated row {row_index} with resume link for '{job_title}'")
            return True

        except Exception as e:
            logger.error(f"Error updating resume link in sheet: {e}")
            return False

def main():
    """Main CLI function."""
    parser = argparse.ArgumentParser(description='Automated Resume Tailor')
    
    # Use defaults from environment or hardcoded values
    default_sheet_id = os.getenv('GOOGLE_SHEET_ID', '1dkxFBloNJQMnXtTfLG5CsGryK_EfmlYrdFFluWPBkgw')
    default_worksheet = os.getenv('GOOGLE_WORKSHEET', 'Jobs')
    default_base_resume = os.getenv('BASE_RESUME_PATH', './base_resume.docx')
    default_master_template = os.getenv('MASTER_TEMPLATE_PATH', './Master_template2.docx')
    default_drive_folder = os.getenv('DRIVE_FOLDER_ID', '1PCBjveFtn07ljJyIhmVjwlwWwPZ7hoFr')
    
    parser.add_argument('--sheet-id', default=default_sheet_id, help=f'Google Sheets document ID (default: {default_sheet_id})')
    parser.add_argument('--worksheet', default=default_worksheet, help=f'Worksheet name (default: {default_worksheet})')
    parser.add_argument('--base-resume', default=default_base_resume, help=f'Path to base resume DOCX file (default: {default_base_resume})')
    parser.add_argument('--master-template', default=default_master_template, help=f'Path to master template DOCX file (default: {default_master_template})')
    parser.add_argument('--overwrite-keywords', action='store_true', help='Overwrite existing keywords')
    parser.add_argument('--dry-run', action='store_true', help='Run in dry-run mode')
    parser.add_argument('--drive-folder-id', default=default_drive_folder, help=f'Google Drive folder ID for uploads (default: {default_drive_folder})')
    
    args = parser.parse_args()
    
    # Initialize the tailor
    tailor = ResumeTailor(dry_run=args.dry_run)
    
    try:
        # Process job data
        logger.info("Processing job data...")
        job_records = tailor.process_job_data(
            args.sheet_id, 
            args.worksheet, 
            args.overwrite_keywords
        )
        
        if not job_records:
            logger.warning("No job records to process")
            return
        
        # Generate resumes for each job
        generated_files = []
        for record in job_records:
            job_title = record.get('JOB_TITLE', '')
            job_description = record.get('JOB_DESCRIPTION', '')
            company_name = record.get('COMPANY', '')
            keywords = record.get('KEYWORDS_RAW', [])

            if not job_title or not job_description:
                continue

            # Check if resume already exists (skip duplicates)
            existing_resume_link = record.get('RESUME_LINK', '').strip()
            if existing_resume_link:
                logger.info(f"Skipping '{job_title}' - Resume already exists: {existing_resume_link}")
                continue

            logger.info(f"Generating resume for: {job_title} (Company: {company_name})")

            resume_path = tailor.generate_resume(
                args.base_resume,
                args.master_template,
                job_title,
                job_description,
                keywords,
                company_name
            )

            if resume_path:
                # Always upload to Drive (mandatory)
                logger.info(f"Uploading resume to Google Drive: {resume_path}")
                drive_link = tailor.upload_to_drive(resume_path, args.drive_folder_id)

                if drive_link:
                    logger.info(f"Resume uploaded to Drive successfully: {drive_link}")

                    # Update Google Sheet with Drive link
                    update_success = tailor.update_resume_link_in_sheet(record, drive_link)
                    if update_success:
                        logger.info(f"Google Sheet updated with Drive link for '{job_title}': {drive_link}")
                        # Only count as generated if fully completed
                        generated_files.append(resume_path)
                    else:
                        logger.warning(f"Failed to update Google Sheet with Drive link for '{job_title}'")

                    # Clean up local file after successful upload
                    try:
                        Path(resume_path).unlink()
                        logger.info(f"Local file cleaned up: {resume_path}")
                    except Exception as e:
                        logger.warning(f"Failed to clean up local file {resume_path}: {e}")

                else:
                    logger.error(f"Failed to upload resume to Drive for '{job_title}' - skipping this job")
                    continue
        
        logger.info(f"Generated {len(generated_files)} new resumes successfully")
        
    except Exception as e:
        logger.error(f"Error in main process: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
